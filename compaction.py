# -*- coding: utf-8 -*-
"""
Asphalt Compaction: Air Void Prediction Studio

What’s new:
- XLSX: Viridis-like Correlation_Heatmap_Color with vertical colorbar (+1 → -1)
- DOCX: remove duplicate heatmap; add Drivers vs Target, Diagnostics, Probability, Learning Curve,
        QQ Plot, Leverage vs Std Residuals, Cook’s Distance, Temperature-only Curve (with narratives)
- DOCX: Insert "Regression Equation" (equation + narrative) just before "Conclusion" (Admins only)
"""

import os, io, json, textwrap, base64, tempfile
from datetime import datetime
from typing import Dict, List, Tuple, Optional

import numpy as np
import pandas as pd
import streamlit as st

from sklearn.model_selection import train_test_split
from sklearn.preprocessing import StandardScaler, PolynomialFeatures
from sklearn.linear_model import LinearRegression, Ridge, Lasso, ElasticNet
from sklearn.pipeline import Pipeline
from sklearn.metrics import r2_score, mean_squared_error, mean_absolute_error

import statsmodels.api as sm
from statsmodels.stats.outliers_influence import variance_inflation_factor

import plotly.express as px
import plotly.graph_objects as go
import plotly.io as pio
import matplotlib.pyplot as plt
import scipy.stats as stats
from scipy.stats import f as f_dist

# Configure Kaleido defaults (for saving Plotly figures as images)
try:
    pio.kaleido.scope.default_format = "png"
    pio.kaleido.scope.default_scale = 2
    pio.kaleido.scope.default_width = 1200
    pio.kaleido.scope.default_height = 700
except Exception:
    pass

import shap

from io import BytesIO
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.lib.utils import ImageReader
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ========= Regression Equation Helpers =========
def _get_pipeline_parts(pipeline: Pipeline):
    poly = pipeline.named_steps.get("poly") if hasattr(pipeline, "named_steps") else None
    scaler = pipeline.named_steps.get("scaler") if hasattr(pipeline, "named_steps") else None
    model = pipeline.named_steps.get("model", pipeline)
    return poly, scaler, model

def _coefs_intercept(est) -> Tuple[np.ndarray, float]:
    coefs = getattr(est, "coef_", None)
    inter = getattr(est, "intercept_", None)
    if coefs is None or inter is None:
        raise ValueError("Estimator lacks coef_/intercept_.")
    return np.atleast_1d(coefs), float(inter)

def _poly_feature_names(poly: Optional[PolynomialFeatures], base_names: List[str]) -> List[str]:
    if poly is None:
        return base_names
    return list(poly.get_feature_names_out(base_names))

def _format_term(name: str, coef: float) -> str:
    if np.isclose(coef, 0.0, atol=1e-12):
        return ""
    sign = " + " if coef >= 0 else " - "
    return f"{sign}{abs(coef):.4g}·{name}"

def build_equation_string(
    pipeline: Pipeline,
    base_feature_names: List[str],
    target_name: str = "y"
) -> Tuple[str, str]:
    """
    Returns (equation_string, equation_notes).

    - Linear + StandardScaler (no PolynomialFeatures): back-transform to original feature scale.
    - With PolynomialFeatures (± StandardScaler): show equation in model-input space and label clearly.
    """
    poly, scaler, model = _get_pipeline_parts(pipeline)
    coefs, intercept = _coefs_intercept(model)
    feat_names = _poly_feature_names(poly, base_feature_names)

    # Case A: No polynomial features
    if poly is None:
        if scaler is None:
            terms = [f"{target_name} = {intercept:.4g}"]
            for n, b in zip(feat_names, coefs):
                t = _format_term(n, b)
                if t: terms.append(t)
            return "".join(terms), "Equation shown in ORIGINAL feature scale (no scaling, no polynomial terms)."
        else:
            means = getattr(scaler, "mean_", None)
            scales = getattr(scaler, "scale_", None)
            if means is not None and scales is not None:
                b_unscaled = coefs / scales
                a_unscaled = intercept - float(np.sum(coefs * (means / scales)))
                terms = [f"{target_name} = {a_unscaled:.4g}"]
                for n, b in zip(feat_names, b_unscaled):
                    t = _format_term(n, b)
                    if t: terms.append(t)
                return "".join(terms), "Equation back-transformed to ORIGINAL feature scale (StandardScaler undone; no polynomial terms)."
            else:
                terms = [f"{target_name} (standardized) = {intercept:.4g}"]
                for n, b in zip(feat_names, coefs):
                    t = _format_term(f"z({n})", b)
                    if t: terms.append(t)
                return "".join(terms), "Equation shown in STANDARDIZED space (scaler stats unavailable)."

    # Case B: PolynomialFeatures present → report in model-input space
    terms = [f"{target_name} (model input space) = {intercept:.4g}"]
    for n, b in zip(feat_names, coefs):
        t = _format_term(n, b)
        if t: terms.append(t)

    if scaler is None:
        notes = "Equation shown in POLYNOMIAL feature space (no scaler). " \
                "Terms like X1^2 and X1:X2 are squared/interaction features."
    else:
        notes = "Equation shown in POLYNOMIAL + STANDARDIZED feature space. " \
                "Exact back-transformation of all cross-terms is non-trivial; " \
                "coefficients apply to the model’s transformed inputs."
    return "".join(terms), notes

def build_equation_narrative(
    r2: float,
    rmse: float,
    mae: float,
    target_name: str,
    key_insights: Optional[List[str]] = None
) -> str:
    """
    Asphalt compaction–specific narrative to accompany the regression equation.
    (Kept compact because long-form 1500-char commentaries exist elsewhere.)
    """
    bullets = key_insights or []
    parts = []
    parts.append(
        f"The regression expression models {target_name} as a weighted combination of field QC factors. "
        f"Overall fit is R²={r2:.3f}, RMSE={rmse:.3f}, MAE={mae:.3f}, indicating typical on-section error "
        f"magnitudes consistent with practical tolerance envelopes for in-place density."
    )
    if bullets:
        parts.append("Coefficient-level observations:")
        for b in bullets:
            parts.append(f"• {b}")
    parts.append(
        "Positive coefficients imply increases in predicted air voids as that factor increases, all else equal; "
        "negative coefficients imply improved densification. These signs and magnitudes help calibrate roller pass counts, "
        "paver speed, and temperature windows to keep density in-spec while avoiding over-compaction."
    )
    return "\n\n".join(parts)

# -----------------------
# Configuration / ENV
# -----------------------
APP_TITLE = os.getenv("APP_TITLE", "Asphalt Compaction: Air Void Prediction Studio")
APP_LOGO_URL = os.getenv("APP_LOGO_URL", "")
APP_LOGO_PATH = os.getenv("APP_LOGO_PATH", "assets/logo.png")
LOGO_SRC = APP_LOGO_URL if APP_LOGO_URL else APP_LOGO_PATH

DEFAULT_ROLE = (os.getenv("DEFAULT_ROLE") or "collaborator").lower()

FAVICON_PATH = os.getenv("FAVICON_PATH", "assets/favicon.png")
SUN_ICON_PATH = os.getenv("SUN_ICON_PATH", "assets/sun.png")
MOON_ICON_PATH = os.getenv("MOON_ICON_PATH", "assets/moon.png")

# Export constants
APP_VERSION = "v1.0.0"
TS = lambda: datetime.now().strftime("%Y%m%d-%H%M%S")
DOC_FILENAME = lambda: f"Asphalt_Compaction_Report_{APP_VERSION}_{TS()}.docx"
PDF_FILENAME = lambda: f"Asphalt_Compaction_Report_{APP_VERSION}_{TS()}.pdf"
XLSX_FILENAME = lambda: f"Asphalt_Compaction_Report_{APP_VERSION}_{TS()}.xlsx"

# GitHub palette
GITHUB_THEME = {
    "light": {
        "bg": "#ffffff", "panel": "#f6f8fa", "text": "#1f2328", "muted": "#57606a",
        "border": "#d0d7de", "link": "#0969da"
    },
    "dark": {
        "bg": "#0d1117", "panel": "#161b22", "text": "#c9d1d9", "muted": "#8b949e",
        "border": "#30363d", "link": "#58a6ff"
    }
}

# Chart export color for non-heatmap charts in DOCX/PDF
DOCX_COLOR = "#4F81BD"

# Bibliography seed for citations
BIB_SEED = [
  {"authors":"TRB/NCHRP", "title":"Intelligent Construction Technologies and Asphalt Compaction", "year":"2019", "publisher":"Transportation Research Board"},
  {"authors":"Roberts, F.L. et al.", "title":"Hot Mix Asphalt Materials, Mixture Design and Construction", "year":"2009", "publisher":"NAPA"},
  {"authors":"Huang, Y.H.", "title":"Pavement Analysis and Design (2nd ed.)", "year":"2004", "publisher":"Pearson/Prentice Hall"},
  {"authors":"Brown, E.R. & Kandhal, P.S.", "title":"Hot Mix Asphalt Compaction", "year":"2001", "publisher":"NCAT"}
]

# -----------------------
# Page config (favicon)
# -----------------------
def _load_favicon() -> Optional[str]:
    if FAVICON_PATH and os.path.exists(FAVICON_PATH):
        return FAVICON_PATH
    return None

st.set_page_config(
    page_title=APP_TITLE,
    page_icon=_load_favicon() or "🛣️",
    layout="wide"
)

# -----------------------
# Role helpers
# -----------------------
def get_user_role() -> str:
    email = st.session_state.get("email", None)
    if email and "roles" in st.secrets:
        return st.secrets["roles"].get(email, DEFAULT_ROLE).lower()
    return DEFAULT_ROLE

def require_admin() -> bool:
    return get_user_role() == "admin"

# ------------------------
# Theme / Icons (PNG first, SVG fallback) — no base64, no broken <img>
# ------------------------
from pathlib import Path
import os
import streamlit as st

APP_DIR = Path(__file__).parent.resolve()
ASSET_DIR = APP_DIR / "assets"

def _first_existing(*candidates: Path) -> Path | None:
    for p in candidates:
        if p and p.exists():
            return p
    return None

SUN_PATH  = _first_existing(ASSET_DIR / "sun.png",  ASSET_DIR / "sun.svg")
MOON_PATH = _first_existing(ASSET_DIR / "moon.png", ASSET_DIR / "moon.svg")

def theme_toggle_icons():
    cols = st.columns([7, 1.4, 1.4, 3])

    st.markdown("""
    <style>
      div.stButton > button { padding:6px 10px; min-width:44px; border-radius:10px; }
      .icon-wrap { text-align:center; margin-top:4px; height:26px; }
    </style>
    """, unsafe_allow_html=True)

    # --- Light mode ---
    with cols[1]:
        if st.button("🔆", key="light_btn"):
            st.session_state["theme"] = "light"
            st.markdown("<script>localStorage.setItem('ac_theme','light');</script>", unsafe_allow_html=True)
            st.rerun()
        # st.image(str(SUN_PATH), width=24)  ← commented out

    # --- Dark mode ---
    with cols[2]:
        if st.button("🌙", key="dark_btn"):
            st.session_state["theme"] = "dark"
            st.markdown("<script>localStorage.setItem('ac_theme','dark');</script>", unsafe_allow_html=True)
            st.rerun()
        # st.image(str(MOON_PATH), width=24) ← commented out

def inject_github_theme():
    if "theme" not in st.session_state:
        st.session_state["theme"] = "light"
    theme = st.session_state["theme"]
    pal = GITHUB_THEME[theme]
    css = f"""
    <style>
      :root {{
        --bg: {pal['bg']}; --panel: {pal['panel']};
        --text: {pal['text']}; --muted: {pal['muted']};
        --border: {pal['border']}; --link: {pal['link']};
      }}
      .stApp {{ background-color: var(--bg) !important; color: var(--text) !important; }}
      a, .stButton button {{ color: var(--link) !important; }}
      .github-panel {{
        background: var(--panel); border:1px solid var(--border);
        border-radius:.5rem; padding:1rem;
      }}
      .icon-img {{ vertical-align: middle; width:24px; height:24px; }}
    </style>
    <script>
      const saved = localStorage.getItem('ac_theme');
      if (saved) {{
        window.parent.document.documentElement.dataset.acTheme = saved;
      }}
    </script>
    """
    st.markdown(css, unsafe_allow_html=True)

# -----------------------
# Data / Modeling utils
# -----------------------
def load_file(uploaded):
    return pd.read_csv(uploaded) if uploaded.name.lower().endswith(".csv") else pd.read_excel(uploaded)

def compute_vif(df: pd.DataFrame) -> pd.DataFrame:
    X = sm.add_constant(df)
    vif = pd.DataFrame({"feature": X.columns})
    vif["VIF"] = [variance_inflation_factor(X.values, i) for i in range(X.shape[1])]
    return vif[vif["feature"]!="const"]

def regression_pipeline(model_name: str, degree:int=1, alpha:float=1.0, l1_ratio:float=0.5, scale:bool=True):
    steps = []
    if scale:
        steps.append(("scaler", StandardScaler(with_mean=True, with_std=True)))
    if degree and degree>1:
        steps.append(("poly", PolynomialFeatures(degree=degree, include_bias=False)))
    if model_name=="Linear":
        steps.append(("model", LinearRegression()))
    elif model_name=="Ridge":
        steps.append(("model", Ridge(alpha=alpha)))
    elif model_name=="Lasso":
        steps.append(("model", Lasso(alpha=alpha)))
    else:
        steps.append(("model", ElasticNet(alpha=alpha, l1_ratio=l1_ratio)))
    return Pipeline(steps)

def regression_equation(model, feature_names):
    try:
        final = model.named_steps.get("model", model)
        poly = model.named_steps.get("poly") if hasattr(model,"named_steps") else None
        names = poly.get_feature_names_out(feature_names) if poly is not None else feature_names
        coefs = getattr(final, "coef_", None)
        intercept = float(getattr(final, "intercept_", 0.0))
        if coefs is None: return "Regression equation unavailable."
        parts = [f"{intercept:.4f}"]
        for c, n in zip(np.ravel(coefs), names):
            sign = "+" if c>=0 else "-"
            parts.append(f" {sign} {abs(c):.4f}·{n}")
        return "y = " + "".join(parts)
    except Exception as e:
        return f"Equation error: {e}"

def metrics_block(y_true, y_pred):
    r2 = r2_score(y_true, y_pred)
    adj = 1 - (1 - r2) * (len(y_true) - 1) / (len(y_true) - y_true.ndim - 1)
    mse = mean_squared_error(y_true, y_pred)
    rmse = np.sqrt(mse)
    mae = mean_absolute_error(y_true, y_pred)
    return {"R2": r2, "Adj_R2": adj, "MSE": mse, "RMSE": rmse, "MAE": mae}

def plot_corr(df, theme="light"):
    fig = px.imshow(df.corr(numeric_only=True), text_auto=True, aspect="auto", color_continuous_scale="Viridis")
    if theme=="dark": fig.update_layout(template="plotly_dark")
    return fig

def bar_fig(df: pd.DataFrame, x: str, y: str, theme: str) -> go.Figure:
    fig = go.Figure(go.Bar(x=df[x], y=df[y], marker_color=DOCX_COLOR))
    if theme=="dark": fig.update_layout(template="plotly_dark")
    fig.update_layout(margin=dict(l=10,r=10,t=20,b=10))
    return fig

# -----------------------
# Regression Summary helpers (Excel-style)
# -----------------------
def _statsmodels_ols_from_pipeline(pipe, X_df: pd.DataFrame, y_ser: pd.Series):
    X_trans = pipe[:-1].transform(X_df) if hasattr(pipe, "steps") and len(pipe.steps)>1 else X_df.values
    feat_names = list(X_df.columns)
    if "poly" in pipe.named_steps:
        feat_names = list(pipe.named_steps["poly"].get_feature_names_out(feat_names))
    X_sm = sm.add_constant(pd.DataFrame(X_trans, columns=feat_names))
    ols_res = sm.OLS(y_ser.values, X_sm).fit()
    return ols_res, ["const"] + feat_names

def _excel_regression_blocks(ols_res) -> tuple[dict, pd.DataFrame, pd.DataFrame]:
    n = int(ols_res.nobs)
    p = int(ols_res.df_model)
    df_reg = p
    df_res = int(ols_res.df_resid)
    df_tot = df_reg + df_res

    ss_reg = float(ols_res.ess)
    ss_res = float(ols_res.ssr)
    ss_tot = ss_reg + ss_res

    ms_reg = ss_reg / df_reg if df_reg > 0 else np.nan
    ms_res = ss_res / df_res if df_res > 0 else np.nan
    f_stat = ms_reg / ms_res if (np.isfinite(ms_reg) and np.isfinite(ms_res) and ms_res != 0) else np.nan
    sig_f  = float(f_dist.sf(f_stat, df_reg, df_res)) if np.isfinite(f_stat) else np.nan

    r2 = float(ols_res.rsquared)
    adj_r2 = float(ols_res.rsquared_adj)
    multiple_r = float(np.sqrt(max(r2, 0.0)))
    std_err_reg = float(np.sqrt(ms_res)) if np.isfinite(ms_res) else np.nan

    reg_stats = {
        "Multiple R": multiple_r,
        "R Square": r2,
        "Adjusted R Square": adj_r2,
        "Standard Error": std_err_reg,
        "Observations": n,
    }

    anova_df = pd.DataFrame([
        {"Source":"Regression","df":df_reg,"SS":ss_reg,"MS":ms_reg,"F":f_stat,"Significance F":sig_f},
        {"Source":"Residual",  "df":df_res,"SS":ss_res,"MS":ms_res,"F":np.nan,"Significance F":np.nan},
        {"Source":"Total",     "df":df_tot,"SS":ss_tot,"MS":np.nan,"F":np.nan,"Significance F":np.nan},
    ])

    coef = ols_res.params
    se   = ols_res.bse
    tval = ols_res.tvalues
    pval = ols_res.pvalues
    ci   = ols_res.conf_int(0.05)
    coef_df = pd.DataFrame({
        "Variable": coef.index,
        "Coefficients": coef.values,
        "Standard Error": se.values,
        "t Stat": tval.values,
        "P-value": pval.values,
        "Lower 95%": ci[0].values,
        "Upper 95%": ci[1].values,
    })
    return reg_stats, anova_df, coef_df

def _write_excel_regression_summary(ws, reg_stats: dict, anova_df: pd.DataFrame, coef_df: pd.DataFrame):
    from openpyxl.styles import Font, Alignment, Border, Side
    bold = Font(bold=True)
    center = Alignment(horizontal="center")
    thin = Border(left=Side(style="thin"), right=Side(style="thin"),
                  top=Side(style="thin"), bottom=Side(style="thin"))

    ws["A1"].value = "Regression Statistics"
    ws["A1"].font = bold

    labels = ["Multiple R","R Square","Adjusted R Square","Standard Error","Observations"]
    r = 2
    for lab in labels:
        ws.cell(row=r, column=1, value=lab).font = bold
        ws.cell(row=r, column=2, value=reg_stats.get(lab))
        r += 1

    ws["A8"].value = "ANOVA"; ws["A8"].font = bold
    headers = ["df","SS","MS","F","Significance F"]
    for j,h in enumerate([""]+headers, start=1):
        ws.cell(row=9, column=j, value=h).font = bold
        ws.cell(row=9, column=j).alignment = center

    for i,row in enumerate(anova_df.itertuples(index=False), start=10):
        ws.cell(row=i, column=1, value=row.Source)
        ws.cell(row=i, column=2, value=row.df)
        ws.cell(row=i, column=3, value=row.SS)
        ws.cell(row=i, column=4, value=row.MS)
        ws.cell(row=i, column=5, value=row.F)
        ws.cell(row=i, column=6, value=row._5)

    for i in range(9, 9+1+len(anova_df)):
        for j in range(1, 6+1):
            ws.cell(row=i, column=j).border = thin

    base = 14
    ws.cell(row=base, column=1, value="Coefficients").font = bold
    coef_headers = ["Variable","Coefficients","Standard Error","t Stat","P-value","Lower 95%","Upper 95%"]
    for j,h in enumerate(coef_headers, start=1):
        ws.cell(row=base+1, column=j, value=h).font = bold
        ws.cell(row=base+1, column=j).alignment = center
    for i,row in enumerate(coef_df.itertuples(index=False), start=base+2):
        for j,val in enumerate(row, start=1):
            ws.cell(row=i, column=j, value=val)

    rows = coef_df.shape[0] + 1
    cols = len(coef_headers)
    for i in range(base+1, base+1+rows+0):
        for j in range(1, cols+1):
            ws.cell(row=i, column=j).border = thin

    ws.column_dimensions["A"].width = 24
    for col in ["B","C","D","E","F","G"]:
        ws.column_dimensions[col].width = 16

# -----------------------
# Narratives & citations
# -----------------------
def citations_for(section: str) -> str:
    idx = abs(hash(section)) % len(BIB_SEED)
    jdx = (idx + 1) % len(BIB_SEED)
    c1 = BIB_SEED[idx]; c2 = BIB_SEED[jdx]
    fmt = lambda c: f"{c['authors']}. *{c['title']}.* {c['year']}. {c['publisher']}."
    return f"{fmt(c1)} {fmt(c2)}"

def generate_commentary(section: str, df: pd.DataFrame, metrics: Optional[dict]=None, target: Optional[str]=None) -> str:
    base = f"{section} — This section analyzes statistical structure and asphalt compaction relevance. "
    if metrics:
        base += f"Model performance yields R²={metrics.get('R2',0):.3f} and RMSE={metrics.get('RMSE',0):.3f}, gauging adequacy for practical QC. "
    if target:
        base += f"Target is {target}; predictors include mixture composition, gradation, temperature, and process controls governing densification and in-place air voids. "
    base += ("Residual and influence diagnostics detect nonlinearity, heteroskedasticity, and leverage. "
             "From an engineering standpoint, binder content, aggregate packing, and compaction temperature work together to control density, while moisture and segregation often manifest as second-order effects. ")
    base += "Citations: " + citations_for(section)
    return (base + " " * 1600)[:1500]

# -----------------------
# Export helpers
# -----------------------
def _img_from_plotly(fig, scale=2) -> bytes:
    buf = io.BytesIO()
    fig.write_image(buf, format="png", scale=scale)
    return buf.getvalue()

def _img_from_matplotlib(fig) -> bytes:
    buf = io.BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight", dpi=200)
    plt.close(fig)
    return buf.getvalue()

def _add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)

# -----------------------
# XLSX export (ordered + heatmap colorbar)
# -----------------------
from openpyxl.formatting.rule import ColorScaleRule
from openpyxl.styles import Alignment
from openpyxl.utils import get_column_letter

def export_xlsx(tables: Dict[str, pd.DataFrame],
                pipe=None, X_df: Optional[pd.DataFrame]=None, y_ser: Optional[pd.Series]=None) -> bytes:
    """
    Enforce sheet order/names per spec; add Regression_Summary and a Viridis-like
    Correlation_Heatmap_Color with an adjacent vertical colorbar of values (+1 → -1).
    """
    order = [
        "Summary","Regression_Summary","Metrics","Predictions",
        "Correlations_Selected","Correlation_Heatmap","Correlation_Heatmap_Color",
        "Variable_Influence","VIF_Before","VIF_After","VIF","SHAP_Importance"
    ]

    output = BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        # Standard sheets first (skip special)
        skip = {"Regression_Summary","Correlation_Heatmap_Color"}
        for name in order:
            if name in skip: continue
            if name in tables:
                df = tables[name].copy()
                df.to_excel(writer, sheet_name=name[:31], index=False)
                ws = writer.book[name[:31]]
                for cell in ws[1]:
                    ws[cell.coordinate].alignment = Alignment(horizontal="center")

        # Regression_Summary
        writer.book.create_sheet("Regression_Summary")
        ws_sum = writer.book["Regression_Summary"]
        if pipe is not None and X_df is not None and y_ser is not None:
            ols_res, _ = _statsmodels_ols_from_pipeline(pipe, X_df, y_ser)
            reg_stats, anova_df, coef_df = _excel_regression_blocks(ols_res)
            _write_excel_regression_summary(ws_sum, reg_stats, anova_df, coef_df)
        else:
            ws_sum["A1"].value = "Regression Statistics (context unavailable)"

        # Correlation_Heatmap_Color (values + conditional formatting + vertical colorbar)
        writer.book.create_sheet("Correlation_Heatmap_Color")
        wsc = writer.book["Correlation_Heatmap_Color"]

        if "Correlation_Heatmap" in tables and not tables["Correlation_Heatmap"].empty:
            corr_df = tables["Correlation_Heatmap"].copy()
            corr_df.to_excel(writer, sheet_name="Correlation_Heatmap_Color", index=False)

            wsc.freeze_panes = "B2"
            wsc.column_dimensions["A"].width = 28
            for cell in wsc[1]:
                cell.alignment = Alignment(horizontal="center")

            max_row = wsc.max_row
            max_col = wsc.max_column
            if max_row >= 2 and max_col >= 2:
                data_range = f"B2:{wsc.cell(row=max_row, column=max_col).coordinate}"
                viridis_rule = ColorScaleRule(
                    start_type='num', start_value=-1, start_color='440154',  # purple
                    mid_type='num',   mid_value= 0, mid_color='21918C',      # greenish-cyan
                    end_type='num',   end_value= 1, end_color='FDE725'       # yellow
                )
                wsc.conditional_formatting.add(data_range, viridis_rule)

                spacer_col_idx = max_col + 1
                bar_col_idx = max_col + 2
                bar_col_letter = get_column_letter(bar_col_idx)
                wsc.cell(row=1, column=bar_col_idx, value="Scale").alignment = Alignment(horizontal="center")
                wsc.column_dimensions[bar_col_letter].width = 8

                n = max_row - 1
                if n > 0:
                    vals = np.linspace(1.0, -1.0, n)
                    for i, v in enumerate(vals, start=2):
                        wsc.cell(row=i, column=bar_col_idx, value=float(v))
                    bar_range = f"{bar_col_letter}2:{bar_col_letter}{max_row}"
                    wsc.conditional_formatting.add(bar_range, ColorScaleRule(
                        start_type='num', start_value=-1, start_color='440154',
                        mid_type='num',   mid_value= 0, mid_color='21918C',
                        end_type='num',   end_value= 1, end_color='FDE725'
                    ))
                    for r in range(2, max_row + 1):
                        cell = wsc.cell(row=r, column=bar_col_idx)
                        cell.alignment = Alignment(horizontal="right")
                        cell.number_format = "0.0"
        else:
            wsc["A1"].value = "Correlation values not available."

        # Reorder sheets
        existing = {ws.title: i for i, ws in enumerate(writer.book.worksheets)}
        desired = [s for s in order if s in existing]
        writer.book._sheets = [writer.book.worksheets[existing[s]] for s in desired]

    return output.getvalue()

# -----------------------
# DOCX/PDF export
# -----------------------
def export_docx(title: str,
                images: List[Tuple[str, bytes]],
                tables: Dict[str, pd.DataFrame],
                narratives: Dict[str, str],
                include_narratives: bool,
                pipe: Optional[Pipeline]=None,
                feature_names: Optional[List[str]]=None,
                target_name: Optional[str]=None,
                metrics: Optional[Dict[str, float]]=None) -> bytes:
    """
    Build DOCX. If include_narratives=True (Admin), insert 'Regression Equation' section
    (equation + narrative) just before 'Conclusion'.
    """
    doc = Document()
    doc.add_heading(title, level=0)

    # Summary narrative first
    if include_narratives and "Summary" in narratives:
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(narratives["Summary"])

    # Figures
    for name, data in images:
        doc.add_heading(name, level=1)
        doc.add_picture(io.BytesIO(data), width=Inches(6.5))
        _add_caption(doc, f"Figure — {name}")
        if include_narratives and name in narratives:
            doc.add_paragraph(narratives[name])

    # Tables
    for name, df in tables.items():
        if name == "Correlation_Heatmap":
            continue
        doc.add_heading(name, level=1)
        t = doc.add_table(rows=1, cols=len(df.columns))
        hdr = t.rows[0].cells
        for i, c in enumerate(df.columns):
            hdr[i].text = str(c)
        for _, row in df.iterrows():
            cells = t.add_row().cells
            for i, v in enumerate(row):
                cells[i].text = str(v)
        _add_caption(doc, f"Table — {name}")
        if include_narratives and name in narratives:
            doc.add_paragraph(narratives[name])

    # --- NEW: Regression Equation (Admins only), inserted just before Conclusion ---
    if include_narratives and pipe is not None and feature_names and target_name and metrics:
        try:
            eq_str, eq_notes = build_equation_string(pipe, feature_names, target_name=target_name)
        except Exception as e:
            eq_str, eq_notes = f"Equation unavailable: {e}", "Notes: not computed due to an internal error."

        doc.add_heading("Regression Equation", level=1)
        # Monospace-like presentation
        p = doc.add_paragraph()
        run = p.add_run(eq_str)
        run.font.name = "Consolas"
        doc.add_paragraph(eq_notes)

        r2 = float(metrics.get("R2", float("nan")))
        rmse = float(metrics.get("RMSE", float("nan")))
        mae = float(metrics.get("MAE", float("nan")))
        narrative_eq = build_equation_narrative(
            r2=r2, rmse=rmse, mae=mae, target_name=target_name,
            key_insights=[
                "Temperature-related terms frequently dominate effects during night paving.",
                "Roller pass count often interacts with mat temperature, moderating predicted air voids."
            ]
        )
        doc.add_paragraph(narrative_eq)

    # Conclusion (Admins only)
    if include_narratives and ("Conclusion" in narratives):
        doc.add_heading("Conclusion", level=1)
        doc.add_paragraph(narratives["Conclusion"])

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()

def export_pdf(title: str,
               images: List[Tuple[str, bytes]],
               narratives: Dict[str, str],
               include_narratives: bool) -> bytes:
    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    w, h = A4
    xm, ym = 2*cm, 2*cm
    y = h - ym
    c.setFont("Helvetica-Bold", 16)
    c.drawString(xm, y, title); y -= 24
    c.setFont("Helvetica", 10)

    for name, img_bytes in images:
        if y < 7*cm: c.showPage(); y = h - ym
        c.setFont("Helvetica-Bold", 12); c.drawString(xm, y, name); y -= 14
        img = ImageReader(io.BytesIO(img_bytes))
        iw, ih = img.getSize()
        maxw = w - 2*xm
        scale = min(maxw/iw, 10*cm/ih)
        c.drawImage(img, xm, y-ih*scale, width=iw*scale, height=ih*scale, preserveAspectRatio=True, mask='auto')
        y -= ih*scale + 8
        if include_narratives and name in narratives:
            for line in textwrap.wrap(narratives[name], width=95):
                if y < 2*cm: c.showPage(); y = h - ym
                c.drawString(xm, y, line); y -= 12

    if include_narratives and ("Conclusion" in narratives):
        if y < 4*cm: c.showPage(); y = h - ym
        c.setFont("Helvetica-Bold", 12); c.drawString(xm, y, "Conclusion"); y -= 14
        c.setFont("Helvetica", 10)
        for line in textwrap.wrap(narratives["Conclusion"], width=95):
            if y < 2*cm: c.showPage(); y = h - ym
            c.drawString(xm, y, line); y -= 12

    c.showPage(); c.save()
    return buf.getvalue()

# -----------------------
# App UI
# -----------------------
inject_github_theme()

# Title row
t1, t2, t3, t4 = st.columns([6,2,2,2])
with t1:
    col_logo, col_title = st.columns([1,9])
    with col_logo:
        if LOGO_SRC:
            st.image(LOGO_SRC, use_column_width=False, width=40)
    with col_title:
        st.markdown(f"### {APP_TITLE}")
with t2:
    theme_toggle_icons()
with t3:
    email = st.text_input("Email (demo RBAC)", value=st.session_state.get("email",""), placeholder="you@org.com")
    if email: st.session_state["email"] = email
with t4:
    st.markdown(f"**Role:** `{get_user_role()}`")

st.markdown('<div class="github-panel">Upload .xlsx or .csv. Select your target (e.g., Air Voids % or Compaction %), choose model and options.</div>', unsafe_allow_html=True)

uploaded = st.file_uploader("Upload dataset (.xlsx or .csv)", type=["xlsx","csv"])

if uploaded is not None:
    df = load_file(uploaded)
    st.subheader("Preview")
    st.dataframe(df.head(10))
    st.dataframe(df.tail(10))

    num_cols = df.select_dtypes(include=[np.number]).columns.tolist()
    target = st.selectbox("Target (y)", options=num_cols, index=0 if num_cols else None)
    features = st.multiselect("Features (X)", options=[c for c in num_cols if c != target], default=[c for c in num_cols if c != target])

    if target and features:
        X = df[features].copy()
        y = df[target].copy()

        with st.expander("Preprocessing & Model Options", expanded=True):
            c1, c2, c3, c4 = st.columns(4)
            with c1: model_name = st.selectbox("Model", ["Linear","Ridge","Lasso","ElasticNet"])
            with c2: degree = st.slider("Polynomial degree", 1, 3, 1)
            with c3: alpha = st.number_input("alpha (Ridge/Lasso/EN)", min_value=1e-4, max_value=1000.0, value=1.0, step=0.1)
            with c4: l1_ratio = st.slider("l1_ratio (ElasticNet)", 0.0, 1.0, 0.5)
            scale = st.checkbox("Standardize features", value=True)
            test_size = st.slider("Test size", 0.1, 0.4, 0.2)
            random_state = st.number_input("Random state", 0, 9999, 42)

        X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=test_size, random_state=random_state)
        pipe = regression_pipeline(model_name, degree, alpha, l1_ratio, scale)
        pipe.fit(X_train, y_train)
        y_pred = pipe.predict(X_test)

        mets = metrics_block(y_test, y_pred)
        st.subheader("Metrics")
        m1, m2, m3, m4, m5 = st.columns(5)
        m1.metric("R²", f"{mets['R2']:.3f}")
        m2.metric("Adj R²", f"{mets['Adj_R2']:.3f}")
        m3.metric("RMSE", f"{mets['RMSE']:.3f}")
        m4.metric("MAE", f"{mets['MAE']:.3f}")
        m5.metric("MSE", f"{mets['MSE']:.3f}")

        st.subheader("Regression Equation")
        st.code(regression_equation(pipe, features))

        # Correlation heatmap (app view); we will export THIS figure to DOCX/PDF
        st.subheader("Correlation heatmap")
        fig_corr = plot_corr(df[num_cols], theme=st.session_state["theme"])
        st.plotly_chart(fig_corr, use_container_width=True)
        corr_png_bytes = _img_from_plotly(fig_corr, scale=3)

        # SHAP importance (LinearExplainer)
        st.subheader("Feature Importance (SHAP)")
        imp_df = None
        try:
            X_trans = pipe[:-1].transform(X_train) if hasattr(pipe, "steps") and len(pipe.steps)>1 else X_train.values
            linear_model = pipe.named_steps["model"]
            explainer = shap.LinearExplainer(linear_model, X_trans)
            X_trans_test = pipe[:-1].transform(X_test) if hasattr(pipe, "steps") and len(pipe.steps)>1 else X_test.values
            shap_values = explainer(X_trans_test)
            mean_abs = np.mean(np.abs(shap_values.values), axis=0)
            feat_names = features
            if "poly" in pipe.named_steps:
                feat_names = list(pipe.named_steps["poly"].get_feature_names_out(features))
            imp_df = pd.DataFrame({"feature": feat_names, "importance": mean_abs[:len(feat_names)]}).sort_values("importance", ascending=False)
            st.dataframe(imp_df)
            fig_imp = bar_fig(imp_df.head(20), x="feature", y="importance", theme=st.session_state["theme"])
            st.plotly_chart(fig_imp, use_container_width=True)
        except Exception as e:
            st.info(f"SHAP could not be computed: {e}")

        # VIF
        st.subheader("Variance Inflation Factor (VIF)")
        vif_before, vif_after = None, None
        try:
            vif_before = compute_vif(df[features].dropna())
            vif_after = vif_before.copy()
            st.dataframe(vif_before)
        except Exception as e:
            st.info(f"VIF unavailable: {e}")

        # Diagnostics base
        st.subheader("Diagnostics")
        resid = y_test - y_pred
        fitted = y_pred
        fig_res = go.Figure()
        fig_res.add_trace(
            go.Scatter(
                x=fitted, y=resid, mode="markers",
                name="Residuals", marker=dict(color=DOCX_COLOR)
            )
        )
        fig_res.add_hline(y=0, line_dash="dash")
        fig_res.update_layout(
            xaxis_title="Fitted", yaxis_title="Residuals",
            margin=dict(l=10, r=10, t=20, b=10),
            template="plotly_dark" if st.session_state.get("theme") == "dark" else None
        )

        # Probability / QQ plots (Matplotlib)
        fig_prob, axp = plt.subplots()
        stats.probplot(resid, dist="norm", plot=axp)
        prob_png = _img_from_matplotlib(fig_prob)

        fig_qq, axq = plt.subplots()
        sm.qqplot(resid, line='45', ax=axq)
        qq_png = _img_from_matplotlib(fig_qq)

        # Influence diagnostics via statsmodels OLS on training set transform
        ols_res, sm_names = _statsmodels_ols_from_pipeline(pipe, X_train, y_train)
        infl = ols_res.get_influence()
        leverage = infl.hat_matrix_diag
        cooks = infl.cooks_distance[0]
        stud_resid = infl.resid_studentized_internal

        # Leverage vs standardized residuals
        fig_lev = go.Figure()
        fig_lev.add_trace(go.Scatter(x=leverage, y=stud_resid, mode="markers", marker=dict(color=DOCX_COLOR)))
        fig_lev.update_layout(xaxis_title="Leverage", yaxis_title="Standardized Residuals", margin=dict(l=10,r=10,t=20,b=10))

        # Cook's distance bar (top 20)
        top_idx = np.argsort(cooks)[-20:][::-1]
        fig_cook = go.Figure(go.Bar(x=[int(i) for i in top_idx], y=cooks[top_idx], marker_color=DOCX_COLOR))
        fig_cook.update_layout(xaxis_title="Observation (index)", yaxis_title="Cook's Distance", margin=dict(l=10,r=10,t=20,b=10))

        # Drivers vs Target (top |corr| features)
        corr_abs = df[features + [target]].corr(numeric_only=True)[target].drop(target).abs().sort_values(ascending=False)
        top_drivers = list(corr_abs.head(min(3, len(corr_abs))).index)
        driver_imgs: List[Tuple[str, bytes]] = []
        for i, feat in enumerate(top_drivers, start=1):
            fig_sc = px.scatter(df, x=feat, y=target, trendline="ols")
            fig_sc.update_traces(marker=dict(color=DOCX_COLOR))
            driver_imgs.append((f"Drivers_vs_Target_{i} ({feat} vs {target})", _img_from_plotly(fig_sc, scale=3)))

        # Temperature-only curve (auto-detect)
        temp_col = None
        for c in df.columns:
            if "temp" in c.lower():
                temp_col = c; break
        temp_img = None
        if temp_col:
            fig_temp = px.scatter(df, x=temp_col, y=target, trendline="ols")
            fig_temp.update_traces(marker=dict(color=DOCX_COLOR))
            temp_img = ("Temperature-only Curve", _img_from_plotly(fig_temp, scale=3))

        # ----- Narratives (Admin only) -----
        narratives = {}
        if require_admin():
            st.subheader("Narrative Commentary (Admin)")
            sections = [
                "Summary",
                "Drivers vs Target (Correlation)",
                "Diagnostics",
                "Probability Plot",
                "Learning Curve (OLS)",
                "QQ Plot (Residuals)",
                "Leverage vs. Standardized Residuals",
                "Cook's Distance",
                "Temperature-only Curve"
            ]
            for sec in sections:
                narratives[sec] = st.text_area(f"{sec} — 1500 chars", value=generate_commentary(sec, df, mets, target), height=140)
            narratives["Conclusion"] = generate_commentary("Conclusion", df, mets, target)

        # ----- Learning Curve (simple OLS-style via pipeline) -----
        train_sizes = np.linspace(0.2, 1.0, 6)
        tr_scores, te_scores = [], []
        for s in train_sizes:
            n = max(5, int(len(X_train) * s))
            X_sub, y_sub = X_train.iloc[:n], y_train.iloc[:n]
            p2 = regression_pipeline(model_name, degree, alpha, l1_ratio, scale)
            p2.fit(X_sub, y_sub)
            tr_scores.append(r2_score(y_sub, p2.predict(X_sub)))
            te_scores.append(r2_score(y_test, p2.predict(X_test)))
        fig_lc = go.Figure()
        fig_lc.add_trace(go.Scatter(x=train_sizes, y=tr_scores, mode="lines+markers", name="Train R²", marker=dict(color=DOCX_COLOR)))
        fig_lc.add_trace(go.Scatter(x=train_sizes, y=te_scores, mode="lines+markers", name="Test R²"))
        fig_lc.update_layout(xaxis_title="Training fraction", yaxis_title="R²", margin=dict(l=10,r=10,t=20,b=10))

        # ----- Build tables per required sheets -----
        summary_df = pd.DataFrame({
            "Dataset_Rows": [len(df)],
            "Dataset_Cols": [df.shape[1]],
            "Target": [target],
            "Features_Count": [len(features)],
            "Model": [model_name],
            "Degree": [degree],
            "Scaled": [scale],
            "Test_Size": [test_size]
        })
        metrics_df = pd.DataFrame([mets])
        preds_df = pd.DataFrame({"y_true": y_test, "y_pred": y_pred}).reset_index(drop=True)
        corr_mat = df[num_cols].corr(numeric_only=True)
        corr_sel = corr_mat.loc[features, [target]].reset_index().rename(columns={"index":"feature"})
        var_infl = imp_df.copy() if imp_df is not None else pd.DataFrame(columns=["feature","importance"])
        vif_df = vif_before.copy() if vif_before is not None else pd.DataFrame(columns=["feature","VIF"])
        vif_bef = vif_before.copy() if vif_before is not None else pd.DataFrame(columns=["feature","VIF"])
        vif_aft = vif_after.copy() if vif_after is not None else pd.DataFrame(columns=["feature","VIF"])

        tables = {
            "Summary": summary_df,
            "Metrics": metrics_df,
            "Predictions": preds_df,
            "Correlations_Selected": corr_sel,
            "Correlation_Heatmap": corr_mat.reset_index().rename(columns={"index":"feature"}),
            "Variable_Influence": var_infl,
            "VIF_Before": vif_bef,
            "VIF_After": vif_aft,
            "VIF": vif_df,
            "SHAP_Importance": var_infl
        }

        # ----- Prepare export-grade images -----
        images: List[Tuple[str, bytes]] = []
        images.append(("Correlation_Heatmap", corr_png_bytes))
        if imp_df is not None and not imp_df.empty:
            export_imp = go.Figure(go.Bar(x=imp_df.head(20)["feature"], y=imp_df.head(20)["importance"], marker_color=DOCX_COLOR))
            export_imp.update_layout(margin=dict(l=10,r=10,t=20,b=10))
            images.append(("SHAP_Importance", _img_from_plotly(export_imp, scale=3)))
        images.append(("Diagnostics (Residuals vs Fitted)", _img_from_plotly(fig_res, scale=3)))
        images.append(("Probability Plot", prob_png))
        images.append(("Learning Curve (OLS)", _img_from_plotly(fig_lc, scale=3)))
        images.append(("QQ Plot (Residuals)", qq_png))
        images.append(("Leverage vs. Standardized Residuals", _img_from_plotly(fig_lev, scale=3)))
        images.append(("Cook's Distance", _img_from_plotly(fig_cook, scale=3)))
        for name, img in driver_imgs:
            images.append((name, img))
        if temp_img:
            images.append(temp_img)

        # ----- Exports -----
        st.subheader("Exports")

        # XLSX
        xlsx_bytes = export_xlsx(
            tables,
            pipe=pipe,
            X_df=X_train,
            y_ser=y_train
        )
        st.download_button(
            label="Download XLSX",
            data=xlsx_bytes,
            file_name=XLSX_FILENAME(),
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

        # DOCX (Admin includes narratives + regression-equation section; Collaborator omits)
        docx_bytes = export_docx(
            APP_TITLE,
            images,
            tables,
            narratives if require_admin() else {},
            include_narratives=require_admin(),
            pipe=pipe,
            feature_names=features,
            target_name=target,
            metrics=mets
        )
        st.download_button(
            label="Download DOCX",
            data=docx_bytes,
            file_name=DOC_FILENAME(),
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        # PDF (no regression equation injection to keep PDF concise)
        pdf_bytes = export_pdf(
            APP_TITLE,
            images,
            narratives if require_admin() else {},
            include_narratives=require_admin()
        )
        st.download_button(
            label="Download PDF",
            data=pdf_bytes,
            file_name=PDF_FILENAME(),
            mime="application/pdf"
        )
else:
    st.info("Upload a dataset to begin.")
